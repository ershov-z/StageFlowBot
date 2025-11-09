# core/parser.py
from __future__ import annotations
from docx import Document
from pathlib import Path
import re
import json
from typing import List, Dict, Optional, Tuple
from loguru import logger
from core.types import Actor, Block, Program

# ============================================================
# üé≠ –ó–∞–≥—Ä—É–∑–∫–∞ —Å–ø–∏—Å–∫–∞ –∞–∫—Ç—ë—Ä–æ–≤
# ============================================================

def _load_actor_names() -> set[str]:
    search_paths = [
        Path("/workspace/actors_list.json"),
        Path(__file__).resolve().parents[2] / "actors_list.json",
        Path(__file__).resolve().parents[1] / "actors_list.json",
        Path(__file__).resolve().parent / "actors_list.json",
        Path(__file__).resolve().parent / "data" / "actors_list.json",
    ]
    for path in search_paths:
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    names = {x.strip().lower() for x in json.load(f) if x.strip()}
                    logger.info(f"üé≠ –ó–∞–≥—Ä—É–∂–µ–Ω–æ –∞–∫—Ç—ë—Ä–æ–≤: {len(names)} –∏–∑ {path}")
                    return names
            except Exception as e:
                logger.warning(f"‚ö† –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è {path}: {e}")
    logger.warning("‚ö† actors_list.json –Ω–µ –Ω–∞–π–¥–µ–Ω ‚Äî fallback –∫ –±–∞–∑–æ–≤–æ–º—É –ø–∞—Ä—Å–∏–Ω–≥—É.")
    return set()

ACTOR_NAMES = _load_actor_names()

# ============================================================
# üß© –í—Å–ø–æ–º–æ–≥–∞—Ç–µ–ª—å–Ω—ã–µ —Ñ—É–Ω–∫—Ü–∏–∏ –ø–∞—Ä—Å–∏–Ω–≥–∞
# ============================================================

_SPLIT_RE = re.compile(r"[\n\r\u000b\u2028\u2029;,/\\]+")

def _split(blob: str) -> List[str]:
    if not blob:
        return []
    return [t.strip() for t in _SPLIT_RE.split(blob) if t.strip()]


def _clean_name(token: str) -> str:
    """–£–¥–∞–ª—è–µ—Ç –∑–Ω–∞–∫–∏ %, !, —Ü–∏—Ñ—Ä—ã, –Ω–µ–≤–∏–¥–∏–º—ã–µ —Å–∏–º–≤–æ–ª—ã –∏ –æ–±—Ä–µ–∑–∞–µ—Ç –ø—Ä–æ–±–µ–ª—ã."""
    token = re.sub(r"[%!\d.,]+", "", token)
    token = token.replace("\u200b", "").replace("\xa0", " ")
    return token.strip()


def _try_split_concatenated(token: str) -> List[str]:
    """
    –ü—ã—Ç–∞–µ—Ç—Å—è —Ä–∞–∑–±–∏—Ç—å —Å–ª–µ–ø–ª–µ–Ω–Ω—ã–µ –∏–º–µ–Ω–∞ –Ω–∞ –æ—Å–Ω–æ–≤–µ actors_list.json.
    –ü—Ä–∏–º–µ—Ä: '–ü—É—à–∫–∏–Ω–ò—Å–∞–µ–≤' ‚Üí ['–ü—É—à–∫–∏–Ω', '–ò—Å–∞–µ–≤']
    """
    if not ACTOR_NAMES:
        return [token.strip()]

    clean = token.strip().replace("\u200b", "").replace("\xa0", " ")
    low = clean.lower()
    out, i = [], 0
    names_sorted = sorted(ACTOR_NAMES, key=len, reverse=True)

    while i < len(low):
        matched = False
        for name in names_sorted:
            if low.startswith(name, i):
                out.append(name)
                i += len(name)
                matched = True
                break
        if not matched:
            i += 1

    if len(out) > 1:
        return [s.capitalize() for s in out]
    return [clean.capitalize()]


# ============================================================
# üé≠ –ü–∞—Ä—Å–∏–Ω–≥ –∞–∫—Ç—ë—Ä–æ–≤ –∏ —Ç–µ–≥–æ–≤
# ============================================================

def _parse_actor_tokens(raw: str) -> List[Actor]:
    res: List[Actor] = []
    for tok in _split(raw):
        if not tok:
            continue

        tags = []
        name = tok

        if "%" in name:
            tags.append("later")
        if "!" in name:
            tags.append("early")

        # –Ω–æ–≤—ã–π —Ç–µ–≥ –∑–∫ ‚Üí vo
        if re.search(r"\(?\b–∑\s*–∫\b\)?", name, flags=re.IGNORECASE):
            tags.append("vo")

        # –æ—á–∏—â–∞–µ–º –∏–º—è –æ—Ç —Å–ª—É–∂–µ–±–Ω—ã—Ö –ø–æ–º–µ—Ç–æ–∫ (–≥–∫ –∏ –∑–∫)
        name = re.sub(r"\(?\b(–∑\s*–∫|–≥\s*–∫)\b\)?", "", name, flags=re.IGNORECASE)
        name = _clean_name(name)

        for nm in _try_split_concatenated(name):
            nm = " ".join(nm.split())
            if nm:
                res.append(Actor(name=nm, tags=sorted(set(tags))))
    return res


def _merge_actors(main_list: List[Actor], pp_list: List[Actor]) -> List[Actor]:
    merged = {}
    for a in main_list:
        merged.setdefault(a.name, set()).update(a.tags)
    for a in pp_list:
        merged.setdefault(a.name, set()).update(a.tags)
    return [Actor(name=k, tags=sorted(v)) for k, v in merged.items()]


def _detect_type(title: str) -> str:
    t = (title or "").lower()
    if "[filler]" in t or "—Ç—è–Ω—É—á" in t:
        return "filler"
    if "–ø—Ä–µ–¥–∫—É–ª–∏—Å—å–µ" in t:
        return "prelude"
    if "—Å–ø–æ–Ω—Å–æ—Ä" in t or "sponsor" in t:
        return "sponsor"
    return "performance"


def _is_kv(cell_text: str) -> bool:
    return bool(re.search(r"\b–∫–≤\b", cell_text or "", flags=re.IGNORECASE))


# ============================================================
# üóÇÔ∏è –û–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ —Å—Ö–µ–º—ã –∫–æ–ª–æ–Ω–æ–∫
# ============================================================

def _normalize_header(s: str) -> str:
    return (s or "").strip().lower()


def _guess_mapping_by_header(header_cells: List[str]) -> Optional[Dict[str, int]]:
    h = [_normalize_header(x) for x in header_cells]
    idx = {name: i for i, name in enumerate(h)}

    def find(*aliases) -> Optional[int]:
        for a in aliases:
            if a in idx:
                return idx[a]
        return None

    title_i = find("–Ω–∞–∑–≤–∞–Ω–∏–µ", "–Ω–æ–º–µ—Ä", "–Ω–∞–∑–≤", "title")
    actors_i = find("–∞–∫—Ç–µ—Ä—ã", "–∞–∫—Ç—ë—Ä—ã", "—É—á–∞—Å—Ç–Ω–∏–∫–∏", "actors")
    pp_i     = find("–ø–ø", "pp")
    hire_i   = find("–Ω–∞–∏–º", "–Ω–∞–π–º", "hire")
    resp_i   = find("–æ—Ç–≤–µ—Ç—Å—Ç–≤–µ–Ω–Ω—ã–π", "–æ—Ç–≤–µ—Ç—Å—Ç–≤", "responsible")
    kv_i     = find("–∫–≤", "kv")
    num_i    = find("‚Ññ", "–Ω–æ–º–µ—Ä", "num", "#", "n")

    if title_i is not None and actors_i is not None and pp_i is not None and kv_i is not None:
        if num_i is None:
            num_i = 0
        if hire_i is None:
            hire_i = 4 if len(h) > 4 else None
        if resp_i is None:
            resp_i = 5 if len(h) > 5 else None
        return {"num": num_i, "title": title_i, "actors": actors_i, "pp": pp_i,
                "hire": hire_i, "resp": resp_i, "kv": kv_i}

    if actors_i is not None and pp_i is not None and kv_i is not None and title_i is None:
        num_i = num_i if num_i is not None else 0
        return {"num": num_i, "title": None, "actors": actors_i, "pp": pp_i,
                "hire": 3 if len(h) > 3 else None,
                "resp": 4 if len(h) > 4 else None,
                "kv": kv_i}

    return None


def _fallback_mapping_by_count(n_cols: int) -> Dict[str, int | None]:
    if n_cols >= 7:
        return {"num": 0, "title": 1, "actors": 2, "pp": 3, "hire": 4, "resp": 5, "kv": 6}
    return {"num": 0, "title": None, "actors": 1, "pp": 2,
            "hire": 3 if n_cols > 3 else None,
            "resp": 4 if n_cols > 4 else None, "kv": 5 if n_cols > 5 else None}


# ============================================================
# üìò –û—Å–Ω–æ–≤–Ω–æ–π –ø–∞—Ä—Å–µ—Ä
# ============================================================

def parse_docx(path: str) -> Program:
    logger.info(f"üìÑ –ß—Ç–µ–Ω–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {path}")
    doc = Document(path)
    if not doc.tables:
        logger.error("‚ùå –í –¥–æ–∫—É–º–µ–Ω—Ç–µ –Ω–µ—Ç —Ç–∞–±–ª–∏—Ü.")
        return Program(blocks=[])

    table = doc.tables[0]
    rows = table.rows
    if not rows:
        logger.error("‚ùå –ü–µ—Ä–≤–∞—è —Ç–∞–±–ª–∏—Ü–∞ –ø—É—Å—Ç–∞.")
        return Program(blocks=[])

    header_cells = [c.text for c in rows[0].cells]
    mapping = _guess_mapping_by_header(header_cells)
    if mapping is None:
        mapping = _fallback_mapping_by_count(len(rows[0].cells))
        logger.warning("‚ö† –ó–∞–≥–æ–ª–æ–≤–∫–∏ –Ω–µ —Ä–∞—Å–ø–æ–∑–Ω–∞–Ω—ã ‚Äî –∏—Å–ø–æ–ª—å–∑—É–µ–º —ç–≤—Ä–∏—Å—Ç–∏–∫—É –ø–æ –ø–æ–∑–∏—Ü–∏—è–º.")

    def get(cells: List[str], key: str) -> str:
        i = mapping.get(key)
        return cells[i].strip() if (i is not None and i < len(cells)) else ""

    blocks: List[Block] = []
    next_id = 1

    for row in rows[1:]:
        cells = [c.text or "" for c in row.cells]
        if not any(x.strip() for x in cells):
            continue

        num_raw = get(cells, "num")
        title   = get(cells, "title")
        actors_raw = get(cells, "actors")
        pp_raw     = get(cells, "pp")
        hire       = get(cells, "hire")
        resp       = get(cells, "resp")
        kv_raw     = get(cells, "kv")

        if not title:
            maybe_title = actors_raw.strip()
            lowered = maybe_title.lower()
            if any(x in lowered for x in ("[filler]", "—Ç—è–Ω—É—á", "–ø—Ä–µ–¥–∫—É–ª–∏—Å—å–µ", "—Å–ø–æ–Ω—Å–æ—Ä", "sponsor")):
                title = maybe_title

        main_actors = _parse_actor_tokens(actors_raw)
        pp_actors   = _parse_actor_tokens(pp_raw)
        actors      = _merge_actors(main_actors, pp_actors)

        block_type = _detect_type(title or actors_raw)
        kv = _is_kv(kv_raw)

        blocks.append(Block(
            id=next_id,
            name=title or f"–ë–ª–æ–∫ {next_id}",
            type=block_type,
            actors=actors,
            kv=kv,
            fixed=(block_type in {"prelude", "sponsor"}),
            num=num_raw or "",
            actors_raw=actors_raw or "",
            pp_raw=pp_raw or "",
            hire=hire or "",
            responsible=resp or "",
            kv_raw=kv_raw or "",
        ))
        next_id += 1

    # ============================================================
    # üîí –§–∏–∫—Å–∞—Ü–∏—è –±–ª–æ–∫–æ–≤ –ø–æ –æ–±–Ω–æ–≤–ª—ë–Ω–Ω—ã–º —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è–º
    # ============================================================

    perf_indices = [i for i, b in enumerate(blocks) if b.type == "performance"]

    # —Ñ–∏–∫—Å–∏—Ä—É–µ–º –ø—Ä–µ–¥–∫—É–ª–∏—Å—å–µ –∏ —Å–ø–æ–Ω—Å–æ—Ä–æ–≤
    for b in blocks:
        if b.type in {"prelude", "sponsor"}:
            b.fixed = True

    # —Ñ–∏–∫—Å–∏—Ä—É–µ–º –ø–µ—Ä–≤—ã–µ –¥–≤–∞ –∏ –ø–æ—Å–ª–µ–¥–Ω–∏–µ —á–µ—Ç—ã—Ä–µ –Ω–æ–º–µ—Ä–∞
    for i in range(len(blocks)):
        if i in perf_indices[:2] or i in perf_indices[-4:]:
            blocks[i].fixed = True

    # —Ñ–∏–∫—Å–∏—Ä—É–µ–º —Ç—è–Ω—É—á–∫–∏ –º–µ–∂–¥—É —Ñ–∏–∫—Å–∏—Ä–æ–≤–∞–Ω–Ω—ã–º–∏ –Ω–æ–º–µ—Ä–∞–º–∏
    for i, b in enumerate(blocks):
        if b.type == "filler":
            prev_fixed = i > 0 and blocks[i - 1].fixed
            next_fixed = i < len(blocks) - 1 and blocks[i + 1].fixed
            if prev_fixed and next_fixed:
                b.fixed = True

    logger.info(f"‚úÖ –ü—Ä–æ—á–∏—Ç–∞–Ω–æ –±–ª–æ–∫–æ–≤: {len(blocks)} | performance={len(perf_indices)}")
    return Program(blocks=blocks)
