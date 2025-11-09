from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Dict, List, Optional
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table

from core.types import Arrangement

log = logging.getLogger("stageflow.exporter")

# ============================================================
# Табличные утилиты для сохранения исходного шаблона
# ============================================================

def _normalize_header(value: str) -> str:
    return (value or "").strip().lower()


def _guess_mapping_by_header(header_cells: List[str]) -> Optional[Dict[str, Optional[int]]]:
    normalized = [_normalize_header(cell) for cell in header_cells]
    index = {cell: i for i, cell in enumerate(normalized)}

    def find(*aliases: str) -> Optional[int]:
        for alias in aliases:
            if alias in index:
                return index[alias]
        return None

    num_i = find("№", "номер", "num", "#", "n")
    title_i = find("название", "название номера", "назв", "title")
    actors_i = find("актеры", "актёры", "участники", "actors")
    pp_i = find("пп", "pp")
    hire_i = find("найм", "наим", "hire")
    resp_i = find("ответственный", "ответств", "responsible")
    kv_i = find("кв", "kv")

    if actors_i is None or pp_i is None or kv_i is None:
        return None

    if title_i is None and len(normalized) > 1:
        title_i = 1

    return {
        "num": num_i if num_i is not None else 0,
        "title": title_i,
        "actors": actors_i,
        "pp": pp_i,
        "hire": hire_i,
        "resp": resp_i,
        "kv": kv_i,
    }


def _fallback_mapping_by_count(n_cols: int) -> Dict[str, Optional[int]]:
    if n_cols <= 0:
        return {"num": None, "title": None, "actors": None, "pp": None, "hire": None, "resp": None, "kv": None}

    mapping: Dict[str, Optional[int]] = {
        "num": 0,
        "title": None,
        "actors": None,
        "pp": None,
        "hire": None,
        "resp": None,
        "kv": None,
    }

    next_idx = 1

    if n_cols > next_idx:
        mapping["title"] = next_idx
        next_idx += 1

    if n_cols > next_idx:
        mapping["actors"] = next_idx
        next_idx += 1

    if n_cols > next_idx:
        mapping["pp"] = next_idx
        next_idx += 1

    if n_cols > next_idx:
        mapping["hire"] = next_idx
        next_idx += 1

    if n_cols > next_idx:
        mapping["resp"] = next_idx
        next_idx += 1

    if n_cols > next_idx:
        mapping["kv"] = next_idx

    return mapping


def _clear_table_data(table) -> None:
    while len(table.rows) > 1:
        table._element.remove(table.rows[1]._element)


def _set_cell_text(row, mapping: Dict[str, Optional[int]], key: str, value: str) -> None:
    idx = mapping.get(key)
    if idx is None:
        return
    cells = row.cells
    if idx >= len(cells):
        return
    cells[idx].text = value or ""


def _populate_table(table, blocks, mapping: Dict[str, Optional[int]]) -> None:
    seq = 0
    for block in blocks:
        row = table.add_row()
        if block.type == "performance":
            seq += 1
            number_text = str(seq)
        else:
            number_text = ""

        _set_cell_text(row, mapping, "num", number_text)
        _set_cell_text(row, mapping, "title", block.name or "")
        _set_cell_text(row, mapping, "actors", getattr(block, "actors_raw", "") or "")
        _set_cell_text(row, mapping, "pp", getattr(block, "pp_raw", "") or "")
        _set_cell_text(row, mapping, "hire", getattr(block, "hire", "") or "")
        _set_cell_text(row, mapping, "resp", getattr(block, "responsible", "") or "")

        kv_value = getattr(block, "kv_raw", "") or ("кв" if getattr(block, "kv", False) else "")
        _set_cell_text(row, mapping, "kv", kv_value)


def _score_mapping(mapping: Dict[str, Optional[int]]) -> tuple[int, int]:
    """Оценка соответствия таблицы ожидаемой структуре."""
    required = ("actors", "pp", "kv")
    required_hit = sum(1 for key in required if mapping.get(key) is not None)
    total_hit = sum(1 for value in mapping.values() if value is not None)
    return required_hit, total_hit


def _find_target_table(doc: DocumentType) -> Optional[tuple[Table, Dict[str, Optional[int]]]]:
    """Находит таблицу с наилучшим соответствием ожидаемым колонкам."""
    best: Optional[tuple[Table, Dict[str, Optional[int]], tuple[int, int]]] = None

    for table in doc.tables:
        if not table.rows:
            continue

        header_cells = [cell.text for cell in table.rows[0].cells]
        if not header_cells:
            continue

        mapping = _guess_mapping_by_header(header_cells)
        if mapping is None:
            mapping = _fallback_mapping_by_count(len(header_cells))

        score = _score_mapping(mapping)
        if best is None or score > best[2]:
            best = (table, mapping, score)
            if score[0] == 3:  # нашли таблицу с ключевыми колонками
                break

    if best is None:
        return None

    return best[0], best[1]

# ============================================================
# DOCX и JSON экспорт
# ============================================================

def create_docx(
    blocks,
    path: Path,
    title: str = "StageFlow Program",
    template_path: Path | str | None = None,
):
    """Генерация DOCX-файла программы с сохранением структуры исходного шаблона."""

    template = Path(template_path) if template_path else None
    doc: DocumentType | None = None

    if template and template.exists():
        try:
            doc = Document(template)
            log.debug(f"[EXPORT] Используется шаблон DOCX: {template}")
        except Exception as exc:  # pragma: no cover - защитное логирование
            log.warning(
                f"[EXPORT] Не удалось открыть шаблон {template}: {exc}. Используем пустой документ."
            )
            doc = None

    if doc is None:
        doc = Document()
        doc.add_heading(title, level=1)
        seq = 0
        for block in blocks:
            prefix = ""
            if block.type == "performance":
                seq += 1
                prefix = f"{seq}. "
            doc.add_paragraph(f"{prefix}{block.name or ''}")

        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        log.info(f"[EXPORT] DOCX сохранён (без шаблона): {path.name}")
        return

    table_and_mapping = _find_target_table(doc)
    if table_and_mapping is None:
        log.warning("[EXPORT] Шаблон не содержит подходящих таблиц. Переходим к простому экспорту.")
        create_docx(blocks, path, title=title, template_path=None)
        return

    table, mapping = table_and_mapping

    _clear_table_data(table)
    _populate_table(table, blocks, mapping)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    log.info(f"[EXPORT] DOCX сохранён: {path.name}")


def export_json(arr: Arrangement, path: Path, extra: dict | None = None):
    """Сохраняет JSON-файл с программой."""
    data = {
        "seed": arr.seed,
        "strong_conflicts": arr.strong_conflicts,
        "weak_conflicts": arr.weak_conflicts,
        "fillers_used": arr.fillers_used,
        "meta": arr.meta or {},
        "blocks": [b.to_dict() for b in arr.blocks],
    }
    if extra:
        data.update(extra)

    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    log.info(f"[EXPORT] JSON сохранён: {path.name}")


# ============================================================
# Экспорт одного варианта
# ============================================================

def export_variant(
    arr: Arrangement,
    index: int,
    output_dir: Path,
    template_path: Path | str | None = None,
) -> tuple[str, str]:
    """Экспортирует один вариант в DOCX и JSON."""
    seed = arr.seed
    status = str(arr.meta.get("status", "")).lower()
    is_ideal = status == "ideal"

    label = f"StageFlow_Variant_{index}_seed{seed}"
    if is_ideal:
        label += "_IDEAL"

    docx_path = output_dir / f"{label}.docx"
    json_path = output_dir / f"{label}.json"

    create_docx(
        arr.blocks,
        docx_path,
        title=f"Вариант #{index} (seed={seed}){' — ИДЕАЛЬНЫЙ' if is_ideal else ''}",
        template_path=template_path,
    )
    export_json(arr, json_path, extra={"is_ideal": is_ideal})

    log.info(f"[EXPORT] {'Идеальный ' if is_ideal else ''}вариант seed={seed} → {docx_path.name}")
    return str(docx_path), str(json_path)


# ============================================================
# Пакетный экспорт всех вариантов
# ============================================================

def export_all_variants(
    arrangements: list[Arrangement],
    output_dir: Path,
    template_path: Path | str | None = None,
) -> Path:
    """
    Экспортирует все варианты в DOCX и JSON и собирает ZIP.
    Идеальные варианты помещаются в начало архива.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    log.info("[EXPORT_ALL] Начинается пакетный экспорт всех вариантов")

    # Идеалы первыми
    sorted_arrs = sorted(
        arrangements,
        key=lambda a: 0 if str(a.meta.get("status", "")).lower() == "ideal" else 1
    )

    exported_files: list[str] = []
    for idx, arr in enumerate(sorted_arrs, start=1):
        docx_path, json_path = export_variant(arr, idx, output_dir, template_path=template_path)
        exported_files.extend([docx_path, json_path])

    zip_path = output_dir / "StageFlow_Results.zip"
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as zipf:
        for fpath in exported_files:
            arcname = Path(fpath).name
            zipf.write(fpath, arcname)
            log.info(f"[EXPORT_ALL] Добавлен в архив: {arcname}")

    log.info(f"[EXPORT_ALL] Архив готов: {zip_path}")
    return zip_path


# ============================================================
# Совместимость со старым API
# ============================================================

export_all = export_all_variants


# ============================================================
# Для ручного теста
# ============================================================

if __name__ == "__main__":
    from core.types import Actor, Block
    test_block = Block(1, "Номер 1", "performance", [Actor("Пушкин")])
    arr = Arrangement(
        seed=0,
        blocks=[test_block],
        fillers_used=0,
        strong_conflicts=0,
        weak_conflicts=0,
        meta={"status": "IDEAL"}
    )
    out_dir = Path("./test_export")
    export_all_variants([arr], out_dir)
