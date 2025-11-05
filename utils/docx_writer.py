# utils/docx_writer.py

from __future__ import annotations

import os
import copy
from pathlib import Path
from typing import Optional, List

from docx import Document
from loguru import logger


# ============================================================
# 🔎 Поиск исходного шаблона (оригинального .docx)
# ============================================================

def _autodetect_template_path(explicit_path: Optional[str | os.PathLike]) -> Path:
    """
    Логика выбора исходного .docx, если template_path не передан:
      1) Если передан явный путь — используем его.
      2) Иначе ищем самый свежий пользовательский .docx в папке ./data,
         исключая файлы, начинающиеся на: output_, parsed_, result_.
      3) Если не нашли — поднимаем ошибку.
    """
    if explicit_path:
        p = Path(explicit_path)
        if not p.exists():
            raise FileNotFoundError(f"Шаблон не найден: {p}")
        return p

    data_dir = Path("data")
    if not data_dir.exists():
        raise FileNotFoundError("Папка data не найдена для автопоиска шаблона.")

    candidates: List[Path] = []
    for p in data_dir.glob("*.docx"):
        name = p.name.lower()
        if name.startswith(("output_", "parsed_", "result_")):
            continue
        candidates.append(p)

    if not candidates:
        raise FileNotFoundError("Не найден ни один пользовательский .docx в папке data.")

    # сортируем по времени изменения (свежий первым)
    candidates.sort(key=lambda x: x.stat().st_mtime, reverse=True)
    chosen = candidates[0]
    logger.info(f"🧭 Автовыбран шаблон: {chosen}")
    return chosen


# ============================================================
# 📑 Поиск нужной таблицы и строк
# ============================================================

def _find_main_table(doc: Document):
    """
    Находит основную таблицу по заголовкам (гибко).
    Ориентиры: наличие столбцов с текстами вроде "Номер", "Актёры", "Ответственный", "ПП", "Найм".
    """
    for table in doc.tables:
        if not table.rows:
            continue
        header_text = " | ".join(c.text.strip().lower() for c in table.rows[0].cells)
        # достаточные признаки "нашей" таблицы
        score = 0
        for token in ("номер", "№", "акт", "пп", "найм", "ответ", "kv", "кв"):
            if token in header_text:
                score += 1
        if score >= 2:
            return table
    # если не нашли по эвристике — как fallback возьмём первую
    if doc.tables:
        logger.warning("Не удалось уверенно определить главную таблицу — берём первую.")
        return doc.tables[0]
    raise ValueError("В документе нет таблиц.")


def _find_row_by_title(table, title: str):
    """
    Ищет строку по названию номера (сравниваем по фрагменту, игнорируя переносы и регистр).
    """
    wanted = (title or "").strip().lower().replace("\n", " ")
    if not wanted:
        return None
    for row in table.rows[1:]:
        for cell in row.cells:
            txt = (cell.text or "").strip().lower().replace("\n", " ")
            if wanted and wanted in txt:
                return row
    return None


def _clone_row(table, source_row):
    """
    Полная копия строки с сохранением форматирования (XML-клонирование).
    """
    new_row = table.add_row()
    new_row._tr = copy.deepcopy(source_row._tr)
    return new_row


# ============================================================
# 🧩 Вставка тянучек
# ============================================================

def _cap_name(name: str) -> str:
    """
    Делает имя с заглавной первой буквой (учитываем, что имена — одиночные слова).
    """
    if not name:
        return ""
    # .capitalize() ок для однословных имён: Пушкин, Исаев, Рожков...
    return name.strip().capitalize()


def _insert_tyanuchka_after(table, prev_row, actor_name: str):
    """
    Вставляет строку "Тянучка" после prev_row, сохраняя стиль.
    Колонки:
      0 — Номер (пусто),
      1 — Название ("Тянучка"),
      2 — Актёры (если актёр ≠ Пушкин/Пятков),
      3 — ПП (если актёр = Пушкин или Пятков),
      остальные — пусто.
    """
    new_row = _clone_row(table, prev_row)
    cells = new_row.cells

    # очистка текста во всех ячейках, стили при этом сохраняются
    for c in cells:
        c.text = ""

    actor_name = _cap_name(actor_name)
    # 1: название
    cells[1].text = "Тянучка"

    # Пушкин/Пятков — в ПП, иначе — в Актёры
    if actor_name in ("Пушкин", "Пятков"):
        cells[3].text = actor_name
    else:
        cells[2].text = actor_name

    return new_row


# ============================================================
# 🧠 Основная функция для бота (совместима с main.py)
# ============================================================

def save_program_to_docx(program_data: list[dict], output_path: str | os.PathLike, template_path: Optional[str | os.PathLike] = None):
    """
    Переставляет строки таблицы по порядку из program_data, добавляет тянучки и перенумеровывает.
    - Если template_path не указан, файл шаблона берётся автоматически:
      самый свежий пользовательский .docx из ./data (не output_/parsed_/result_).
    - Форматирование и всё вне таблицы сохраняется.
    - Нумерация только для "номеров"; тянучки и спонсоры — без номера.
    """
    try:
        tpl_path = _autodetect_template_path(template_path)
        logger.info(f"📝 Формируем итоговый DOCX на базе: {tpl_path}")

        # 1) Загружаем шаблон дважды:
        #    - doc     — сюда собираем результат
        #    - tpl_doc — берём из него исходные строки для клонирования
        doc = Document(tpl_path)
        table = _find_main_table(doc)

        tpl_doc = Document(tpl_path)
        tpl_table = _find_main_table(tpl_doc)

        # 2) Удаляем все строки кроме шапки
        if not table.rows:
            raise ValueError("Таблица пуста.")
        header = table.rows[0]
        old_rows = table.rows[1:]
        for r in old_rows:
            table._tbl.remove(r._tr)

        # 3) Строим в новом порядке
        for item in program_data:
            title = str(item.get("title", "")).strip()
            itype = item.get("type", "") or ""
            if itype != "тянучка":
                # обычный номер/спонсоры — берём строку из шаблона по названию
                src_row = _find_row_by_title(tpl_table, title)
                if src_row is None:
                    logger.warning(f"⚠️ В шаблоне не нашли строку по названию: {title!r}. Пропускаем.")
                    continue
                new_row = _clone_row(table, src_row)
                # прикрепляем клонированную строку (последним действием, чтобы порядок был корректный)
                table._tbl.append(new_row._tr)
            else:
                # тянучка — создаём новую строку на базе предыдущей вставленной
                # определяем ведущего: сначала из actors[0].name, иначе из actors_raw
                actor = ""
                if isinstance(item.get("actors"), list) and item["actors"]:
                    actor = item["actors"][0].get("name", "") or ""
                if not actor:
                    actor = (item.get("actors_raw") or "").strip()
                if not actor:
                    logger.warning("⚠️ Тянучка без актёра — пропускаю вставку.")
                    continue
                prev = table.rows[-1] if len(table.rows) > 1 else header
                _insert_tyanuchka_after(table, prev, actor)

        # 4) Перенумеровываем: только номера (без тянучек и спонсоров)
        logger.info("🔢 Перенумеровываем номера...")
        n = 1
        for row in table.rows[1:]:
            num_cell = row.cells[0]
            title_cell_text = (row.cells[1].text or "").strip().lower()
            # тянучка или спонсоры — без номера
            if title_cell_text.startswith("тянучк") or "спонсор" in title_cell_text:
                num_cell.text = ""
            else:
                num_cell.text = str(n)
                n += 1

        # 5) Сохраняем результат
        outp = Path(output_path)
        outp.parent.mkdir(parents=True, exist_ok=True)
        doc.save(outp)
        logger.success(f"✅ Итоговый DOCX сохранён: {outp.resolve()}")
        return str(outp)

    except Exception as e:
        logger.exception(f"Ошибка при сохранении DOCX: {e}")
        raise
