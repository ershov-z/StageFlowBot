from docx import Document
from loguru import logger
from pathlib import Path
import re
import json

# ============================================================
# 🎭 Загрузка списка актёров из actors_list.json
# ============================================================

ACTOR_NAMES = set()

def _load_actor_names():
    """Пробует найти actors_list.json в нескольких стандартных местах"""
    search_paths = [
        Path(__file__).resolve().parent / "actors_list.json",                     # рядом с файлом
        Path(__file__).resolve().parent / "data" / "actors_list.json",            # utils/data/
        Path(__file__).resolve().parents[1] / "data" / "actors_list.json",        # ../data/
        Path(__file__).resolve().parents[1] / "actors_list.json",                 # в корне проекта
    ]
    for path in search_paths:
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    names = {x.strip().lower() for x in json.load(f) if x.strip()}
                    logger.info(f"🎭 Загружено актёров: {len(names)} из {path}")
                    return names
            except Exception as e:
                logger.warning(f"⚠ Ошибка чтения {path}: {e}")
    logger.warning("⚠ actors_list.json не найден — fallback к базовому парсингу.")
    return set()

ACTOR_NAMES = _load_actor_names()

# ============================================================
# 🧩 Вспомогательные функции
# ============================================================

def _extract_text_with_breaks(cell):
    """Извлекает текст из ячейки Word с сохранением переводов строк"""
    try:
        lines = []
        for p in cell._element.xpath(".//w:p"):
            buf = []
            for r in p.xpath(".//w:r"):
                for t in r.xpath(".//w:t"):
                    if t.text:
                        buf.append(t.text)
                if r.xpath(".//w:br"):
                    buf.append("\n")
            line = "".join(buf).strip()
            if line:
                lines.append(line)
        text = "\n".join(lines).replace("\r", "\n").strip()
        return text
    except Exception as e:
        logger.warning(f"⚠ Ошибка извлечения текста из ячейки: {e}")
        return cell.text.strip() if cell.text else ""


_SPLIT_RE = re.compile(r"[\n\r\u000b\u2028\u2029;,/\\]+")


def _split_people_blob(blob: str) -> list[str]:
    """Разделяет строку актёров на отдельные токены"""
    if not blob:
        return []
    parts = [p.strip() for p in _SPLIT_RE.split(blob) if p.strip()]
    return parts


def _clean_actor_token(token: str) -> str:
    """Удаляет мусорные символы и подготавливает имя"""
    token = re.sub(r"[%!\d.,]+", "", token)  # убираем проценты, восклицательные, цифры
    token = token.strip()
    return token


def _try_split_concatenated(token: str) -> list[str]:
    """
    Пытается разбить склеенные имена (например 'ИланаКсюша')
    по известным именам из ACTOR_NAMES.
    """
    if not ACTOR_NAMES:
        return [token]
    low = token.lower()
    found = []
    i = 0
    while i < len(low):
        match = None
        for name in sorted(ACTOR_NAMES, key=len, reverse=True):
            if low.startswith(name, i):
                found.append(name)
                i += len(name)
                match = True
                break
        if not match:
            i += 1
    if len(found) > 1:
        return [n.capitalize() for n in found]
    return [token]


# ============================================================
# 🎯 Основной парсер актёров
# ============================================================

def parse_actors(raw: str) -> list[dict]:
    """
    Парсит список актёров и теги:
      %  → 'later'
      !  → 'early'
      (гк), гк, Гк, (ГК), г к → 'gk'
    """
    if not raw:
        return []

    result = []
    for token in _split_people_blob(raw):
        if not token.strip():
            continue
        tags = set()
        name = token.strip()

        # Определяем теги
        if "%" in name:
            tags.add("later")
        if "!" in name:
            tags.add("early")

        # Поиск любых вариантов "гк"
        if re.search(r"\(?\bг\s*к\b\)?", name, flags=re.IGNORECASE):
            tags.add("gk")

        # Убираем метки из имени
        name = re.sub(r"\(?\bг\s*к\b\)?", "", name, flags=re.IGNORECASE)
        name = _clean_actor_token(name)

        # Пытаемся разделить склеенные имена
        names = _try_split_concatenated(name)

        for nm in names:
            nm = " ".join(nm.split())
            if nm:
                result.append({"name": nm, "tags": sorted(list(tags))})

    return result


# ============================================================
# 📘 Чтение таблицы и формирование структуры данных
# ============================================================

def read_program(path: str):
    """Читает первую таблицу из DOCX и возвращает структурированные данные"""
    logger.info(f"📄 Чтение документа: {path}")
    doc = Document(path)
    if not doc.tables:
        logger.error("❌ В документе нет таблиц.")
        return []

    table = doc.tables[0]
    rows = table.rows
    if len(rows) < 2:
        logger.error("❌ Таблица пуста.")
        return []

    data = []
    for i, row in enumerate(rows[1:], start=1):
        cells = row.cells
        texts = [_extract_text_with_breaks(c) for c in cells]
        if not any(texts):
            continue

        num = texts[0] if len(texts) > 0 else ""
        title = texts[1] if len(texts) > 1 else ""
        actors_raw = texts[2] if len(texts) > 2 else ""
        pp = texts[3] if len(texts) > 3 else ""  # колонка ПП
        hire = texts[4] if len(texts) > 4 else ""
        responsible = texts[5] if len(texts) > 5 else ""
        kv = "кв" in (texts[6].lower() if len(texts) > 6 and texts[6] else "")

        entry = {
            "order": i,
            "num": num,
            "title": title,
            "actors_raw": actors_raw,
            "pp": pp,
            "hire": hire,
            "responsible": responsible,
            "kv": kv,
            "type": "обычный",
        }

        lower_title = title.lower()
        if "предкулисье" in lower_title:
            entry["type"] = "предкулисье"
        elif "спонсор" in lower_title:
            entry["type"] = "спонсоры"
        elif "тянуч" in lower_title:
            entry["type"] = "тянучка"

        # Объединяем актёров из основной колонки и ПП
        main_actors = parse_actors(actors_raw)
        pp_actors = parse_actors(pp)

        merged_actors = {a["name"]: set(a["tags"]) for a in main_actors}
        for pa in pp_actors:
            name = pa["name"]
            tags = set(pa["tags"])
            if name in merged_actors:
                merged_actors[name].update(tags)
            else:
                merged_actors[name] = tags

        entry["actors"] = [
            {"name": name, "tags": sorted(list(tags))} for name, tags in merged_actors.items()
        ]

        data.append(entry)

    logger.info(f"✅ Прочитано {len(data)} строк.")
    return data


# ============================================================
# 🧪 Тест
# ============================================================

if __name__ == "__main__":
    import json
    test_str = "Илана(гк)! Ксюша Гк% Пушкин (ГК)! Брекоткин г к%%"
    print(json.dumps(parse_actors(test_str), ensure_ascii=False, indent=2))
